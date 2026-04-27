from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from TaoxinAI.repositories.chroma_repo import ChromaRepository
from TaoxinAI.schemas.intake import CaseInput
from TaoxinAI.schemas.reasoning import (
    AnalysisResponse,
    DocumentGenerationRequest,
    DocumentGenerationResponse,
    FullAnalysisResponse,
)
from TaoxinAI.services.fact.rule_based_extractor import RuleBasedFactExtractor
from TaoxinAI.services.intake.entity_extractor import EntityExtractor
from TaoxinAI.services.intake.normalizer import IntakeNormalizer
from TaoxinAI.services.issue.rule_based_identifier import RuleBasedIssueIdentifier
from TaoxinAI.services.matching.element_matcher import ElementMatcher
from TaoxinAI.services.reasoning.irac_generator import IRACGenerator
from TaoxinAI.services.reasoning.reason_builder import ReasonBuilder
from TaoxinAI.services.retrieval.case_retriever import CaseRetriever
from TaoxinAI.services.retrieval.law_retriever import LawRetriever
from TaoxinAI.services.retrieval.reranker import Reranker
from TaoxinAI.services.retrieval.template_retriever import TemplateRetriever
from TaoxinAI.services.verification.completeness_checker import CompletenessChecker


class CaseAnalysisPipeline:
    """High-level workflow for the wage-recovery reasoning prototype."""

    def __init__(self, law_collection: Any, case_collection: Any, llm_client: Any):
        self.llm_client = llm_client
        self.intake_normalizer = IntakeNormalizer()
        self.entity_extractor = EntityExtractor()
        self.law_repo = ChromaRepository(law_collection)
        self.case_repo = ChromaRepository(case_collection)
        template_root = Path(__file__).resolve().parents[1] / "templates" / "taoxin"
        self.law_retriever = LawRetriever(self.law_repo)
        self.case_retriever = CaseRetriever(self.case_repo)
        self.template_retriever = TemplateRetriever(template_root)
        self.reranker = Reranker()
        self.issue_identifier = RuleBasedIssueIdentifier()
        self.fact_extractor = RuleBasedFactExtractor()
        self.element_matcher = ElementMatcher()
        self.reason_builder = ReasonBuilder()
        self.irac_generator = IRACGenerator()
        self.completeness_checker = CompletenessChecker()

    async def analyze_chat_turn(self, case_input: CaseInput) -> AnalysisResponse:
        query_text = self.intake_normalizer.normalize(self._get_last_user_message(case_input))
        retrieved = self._retrieve_authorities(query_text)
        stage = case_input.extracted_info.get("_stage", "initial")
        next_stage, fallback_reply, quick_replies = self._chat_flow(stage)

        extracted_info = dict(case_input.extracted_info)
        extracted_info["_stage"] = next_stage
        extracted_info.update(self.entity_extractor.extract(query_text))

        issues = self.issue_identifier.identify(case_input, retrieved)
        reply, reply_source = await self._generate_chat_reply(
            user_message=query_text,
            current_stage=stage,
            next_stage=next_stage,
            extracted_info=extracted_info,
            retrieved_authorities=retrieved,
            fallback_reply=fallback_reply,
            quick_replies=quick_replies,
        )
        can_generate_doc = bool(
            extracted_info.get("employer_type")
            and any(key in extracted_info for key in ["amount", "evidence_types", "has_contract"])
        )

        return AnalysisResponse(
            reply=reply,
            reply_source=reply_source,
            quick_replies=quick_replies,
            conversation_stage=next_stage,
            extracted_info=extracted_info,
            retrieved_authorities=retrieved,
            issues=issues,
            facts=[],
            element_matches=[],
            reasoning_summary={},
            verification={},
            can_generate_doc=can_generate_doc,
        )

    async def run_full_analysis(self, case_input: CaseInput) -> FullAnalysisResponse:
        query_text = self.intake_normalizer.normalize(self._get_last_user_message(case_input))
        retrieved = self._retrieve_authorities(query_text)
        issues = self.issue_identifier.identify(case_input, retrieved)
        facts = self.fact_extractor.extract(case_input, issues)
        matches = self.element_matcher.match(issues, facts, retrieved)
        reasons = self.reason_builder.build(issues, matches)
        irac = self.irac_generator.generate(issues, matches, reasons, retrieved)
        verification = self.completeness_checker.check(issues, matches, irac)

        return FullAnalysisResponse(
            case_summary=query_text,
            retrieved_authorities=retrieved,
            issues=issues,
            facts=facts,
            element_matches=matches,
            reason_bundles=reasons,
            irac_drafts=irac,
            verification=verification,
        )

    async def generate_document(
        self,
        payload: DocumentGenerationRequest,
    ) -> DocumentGenerationResponse:
        analysis = payload.analysis
        extracted = payload.case_input.extracted_info
        document_type = payload.document_type or self.resolve_document_type(payload.case_input)
        parties = self._resolve_document_parties(document_type)
        amount = extracted.get("amount") or extracted.get("wage_amount") or "待补充"
        employer = extracted.get("employer_name") or extracted.get("employer_type") or "待补充"
        worker = extracted.get("worker_name") or extracted.get("name") or "待补充"
        work_location = extracted.get("location") or extracted.get("work_location") or "待补充"
        wage_period = extracted.get("time_period") or extracted.get("wage_period") or "待补充"
        application_text = (
            analysis.irac_drafts[0].application
            if analysis.irac_drafts
            else "目前已完成基础案件分析，后续将补充完整的 IRAC 论证。"
        )
        citations = self._collect_document_citations(analysis)
        request_items = self._build_request_items(document_type, amount)
        facts_and_reasons = self._build_facts_and_reasons(
            document_type=document_type,
            worker=worker,
            employer=employer,
            amount=amount,
            work_location=work_location,
            wage_period=wage_period,
            application_text=application_text,
            citations=citations,
            extracted=extracted,
        )

        request_text = "\n".join(f"{index + 1}. {item}" for index, item in enumerate(request_items))
        document = (
            f"{self._document_title(document_type)}\n\n"
            f"{parties['applicant_label']}：{worker}\n"
            f"{parties['respondent_label']}：{employer}\n\n"
            "请求事项：\n"
            f"{request_text}\n\n"
            "事实与理由：\n"
            f"{facts_and_reasons}\n"
        )

        risk_notes = list(analysis.verification.consistency_warnings)
        if not risk_notes:
            risk_notes.append("当前文书为骨架版本，仍需补充金额、时间和证据细节。")

        return DocumentGenerationResponse(
            document=document,
            citations=citations,
            risk_notes=risk_notes,
        )

    async def generate_document_docx(self, case_input: CaseInput) -> tuple[bytes, str]:
        document_type = self.resolve_document_type(case_input)
        analysis = await self.run_full_analysis(case_input)
        payload = DocumentGenerationRequest(
            document_type=document_type,
            case_input=case_input,
            analysis=analysis,
        )
        generated = await self.generate_document(payload)
        file_bytes = self._render_document_docx(payload, generated)
        return file_bytes, self._document_filename(document_type)

    def _get_last_user_message(self, case_input: CaseInput) -> str:
        for message in reversed(case_input.messages):
            if message.role == "user":
                return message.content
        return ""

    def _retrieve_authorities(self, query_text: str):
        if not query_text:
            return []

        retrieved = []
        retrieved.extend(self.law_retriever.retrieve(query_text, n_results=4))
        retrieved.extend(self.case_retriever.retrieve(query_text, n_results=4))
        retrieved.extend(self.template_retriever.retrieve(query_text, n_results=2))
        return self.reranker.rerank(retrieved)[:6]

    def resolve_document_type(self, case_input: CaseInput) -> str:
        has_arbitration = case_input.extracted_info.get("has_arbitration")
        if has_arbitration is False:
            return "labor_arbitration_application"
        if has_arbitration is True:
            return "civil_complaint"
        return "labor_complaint"

    def _chat_flow(self, stage: str) -> tuple[str, str, list[str]]:
        if stage == "initial":
            return (
                "basic_facts",
                "请先告诉我几个基础事实：您是给公司干活，还是给包工头或老板干活？有没有签合同？",
                ["给公司干活", "给包工头干活", "有合同", "没有合同"],
            )

        if stage == "basic_facts":
            return (
                "evidence_check",
                "我还需要确认证据情况。您目前有没有工资结算单、聊天记录、考勤记录或者转账记录？",
                ["有聊天记录", "有转账记录", "有考勤记录", "暂时没有"],
            )

        return (
            "analysis_ready",
            "基础信息已经记录。下一步可以进入完整分析，输出争点、要件匹配和文书草稿。",
            ["开始完整分析", "继续补充信息"],
        )

    async def _generate_chat_reply(
        self,
        *,
        user_message: str,
        current_stage: str,
        next_stage: str,
        extracted_info: dict,
        retrieved_authorities: list,
        fallback_reply: str,
        quick_replies: list[str],
    ) -> tuple[str, str]:
        create_completion = self._get_completion_creator()
        if create_completion is None:
            return fallback_reply, "rule_based"

        authority_lines = [
            f"- {item.title}: {item.snippet}"
            for item in retrieved_authorities
            if item.source_type == "norm"
        ]
        authority_block = "\n".join(authority_lines[:3]) or "暂无可用法条"
        extracted_info_json = json.dumps(extracted_info, ensure_ascii=False)
        quick_reply_text = "、".join(quick_replies)

        messages = [
            {
                "role": "system",
                "content": (
                    "你是“讨薪助手”，面向农民工和劳动者进行合规的法律辅助。"
                    "\n请在回复中做到："
                    "\n1. 优先引用 1-2 个最相关的法条，格式像《保障农民工工资支付条例》第三十条。"
                    "\n2. 先用通俗中文给出判断，再结合当前阶段只追问最关键的 1-2 个问题。"
                    "\n3. 不要编造法条条款，不要输出 JSON，不要出现 think 标签。"
                    "\n4. 总体控制在 80-180 字。"
                ),
            },
            {
                "role": "user",
                "content": (
                    f"用户最新提问：{user_message or '（用户暂未额外输入）'}\n"
                    f"当前阶段：{current_stage}\n"
                    f"下一阶段：{next_stage}\n"
                    f"已提取信息：{extracted_info_json}\n"
                    f"候选法条：\n{authority_block}\n"
                    f"如需继续追问，尽量和这些快捷回复保持一致：{quick_reply_text}\n"
                    f"如果模型判断信息仍不足，可参考这句回退话术：{fallback_reply}"
                ),
            },
        ]

        try:
            response = await create_completion(
                model="legalone",
                messages=messages,
                temperature=0.2,
            )
        except Exception:
            return fallback_reply, "rule_based"

        content = self._extract_completion_text(response)
        if not content:
            return fallback_reply, "rule_based"

        return content, "legalone"

    def _get_completion_creator(self):
        chat_api = getattr(self.llm_client, "chat", None)
        completions_api = getattr(chat_api, "completions", None) if chat_api else None
        create_completion = getattr(completions_api, "create", None) if completions_api else None
        return create_completion if callable(create_completion) else None

    def _extract_completion_text(self, response: Any) -> str:
        choices = getattr(response, "choices", None) or []
        if not choices:
            return ""

        message = getattr(choices[0], "message", None)
        content = getattr(message, "content", "") if message else ""
        return content.strip()

    def _document_title(self, document_type: str) -> str:
        return {
            "labor_arbitration_application": "劳动仲裁申请书",
            "civil_complaint": "民事起诉状（追索劳动报酬）",
        }.get(document_type, "劳动维权申请书")

    def _document_filename(self, document_type: str) -> str:
        return f"{self._document_title(document_type)}.docx"

    def _resolve_document_parties(self, document_type: str) -> dict[str, str]:
        if document_type == "civil_complaint":
            return {
                "applicant_label": "原告",
                "respondent_label": "被告",
            }
        return {
            "applicant_label": "申请人",
            "respondent_label": "被申请人",
        }

    def _build_request_items(self, document_type: str, amount: str) -> list[str]:
        if document_type == "labor_arbitration_application":
            return [
                f"请求裁决被申请人支付拖欠工资 {amount}。",
                "请求裁决被申请人承担本案仲裁相关责任。",
            ]

        if document_type == "civil_complaint":
            return [
                f"请求判令被告支付拖欠工资 {amount}。",
                "请求判令被告承担本案诉讼费用。",
            ]

        return [
            f"请求依法支付拖欠工资 {amount}。",
            "请求依法明确相关清偿责任并督促履行。",
        ]

    def _build_facts_and_reasons(
        self,
        *,
        document_type: str,
        worker: str,
        employer: str,
        amount: str,
        work_location: str,
        wage_period: str,
        application_text: str,
        citations: list[str],
        extracted: dict,
    ) -> str:
        contract_status = extracted.get("has_contract")
        contract_clause = "双方签有劳动合同。" if contract_status in {True, "有"} else "目前暂无书面劳动合同。"
        evidence_types = extracted.get("evidence_types") or "待补充"
        legal_basis = f"主要依据包括：{'；'.join(citations)}。" if citations else ""

        if document_type == "labor_arbitration_application":
            return (
                f"申请人 {worker} 在 {work_location} 为 {employer} 提供劳动，欠薪时段为 {wage_period}，"
                f"截至目前仍被拖欠工资 {amount}。{contract_clause} 现有证据包括：{evidence_types}。"
                f"{application_text} {legal_basis}".strip()
            )

        if document_type == "civil_complaint":
            return (
                f"原告 {worker} 在 {work_location} 为 {employer} 提供劳动，欠薪时段为 {wage_period}，"
                f"被拖欠工资 {amount}。{contract_clause} 原告已就争议进行前置维权，现依据相关法律规定提起诉讼。"
                f"{application_text} {legal_basis}".strip()
            )

        return (
            f"{worker} 在 {work_location} 为 {employer} 提供劳动，欠薪时段为 {wage_period}，"
            f"被拖欠工资 {amount}。{contract_clause} 现有证据包括：{evidence_types}。"
            f"{application_text} {legal_basis}".strip()
        )

    def _collect_document_citations(self, analysis: FullAnalysisResponse) -> list[str]:
        if analysis.irac_drafts and analysis.irac_drafts[0].citations:
            return analysis.irac_drafts[0].citations

        citations: list[str] = []
        for authority in analysis.retrieved_authorities:
            if authority.source_type != "norm":
                continue
            if authority.title not in citations:
                citations.append(authority.title)
        return citations[:5]

    def _render_document_docx(
        self,
        payload: DocumentGenerationRequest,
        generated: DocumentGenerationResponse,
    ) -> bytes:
        document = Document()
        document.add_heading(self._document_title(payload.document_type), level=0)

        body_lines = [line.strip() for line in generated.document.splitlines() if line.strip()]
        for index, line in enumerate(body_lines):
            if index == 0 and line == self._document_title(payload.document_type):
                continue
            document.add_paragraph(line)

        if generated.citations:
            document.add_heading("主要法律依据", level=1)
            for citation in generated.citations:
                document.add_paragraph(citation, style="List Bullet")

        if generated.risk_notes:
            document.add_heading("补充提醒", level=1)
            for note in generated.risk_notes:
                document.add_paragraph(note, style="List Bullet")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
